"""
Word文档生成器: LaTeX化学式 → Word OMML公式对象
生成可直接打印的化学试卷 .docx
"""
import re
import html
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree


# ─── LaTeX 常用化学符号映射表 ───
LATEX_CHEM_MAP = {
    # 希腊字母
    r"\alpha": "α", r"\beta": "β", r"\gamma": "γ", r"\delta": "δ",
    r"\epsilon": "ε", r"\zeta": "ζ", r"\eta": "η", r"\theta": "θ",
    r"\lambda": "λ", r"\mu": "μ", r"\nu": "ν", r"\xi": "ξ",
    r"\pi": "π", r"\rho": "ρ", r"\sigma": "σ", r"\tau": "τ",
    r"\phi": "φ", r"\chi": "χ", r"\psi": "ψ", r"\omega": "ω",
    r"\Delta": "Δ", r"\Omega": "Ω", r"\Pi": "Π", r"\Sigma": "Σ",
    r"\Gamma": "Γ", r"\Theta": "Θ", r"\Lambda": "Λ",
    # 化学箭头
    r"\rightarrow": "→", r"\leftarrow": "←",
    r"\rightleftharpoons": "⇌", r"\leftrightharpoons": "⇌",
    r"\longrightarrow": "→", r"\xrightarrow": "→",
    r"\uparrow": "↑", r"\downarrow": "↓",
    # 特殊符号
    r"\triangle": "△", r"\Delta": "Δ",
    r"\infty": "∞", r"\pm": "±", r"\mp": "∓", r"\cdot": "·",
    r"\times": "×", r"\div": "÷", r"\approx": "≈",
    r"\geq": "≥", r"\leq": "≤", r"\neq": "≠",
    r"\equiv": "≡", r"\propto": "∝", r"\sim": "∼",
    # 摄氏度等
    r"^\circ C": "°C", r"^\circ": "°",
    # 气体/沉淀箭头
    r"\uparrow": "↑", r"\downarrow": "↓",
}


def _latex_to_unicode(text: str) -> str:
    """将LaTeX化学符号转为Unicode（在Word中直接显示）"""
    for latex, unicode_char in LATEX_CHEM_MAP.items():
        text = text.replace(latex, unicode_char)
    return text


def _parse_latex_subscript_superscript(text: str) -> str:
    """
    解析化学式中的上下标: H_2SO_4 → H₂SO₄, Fe^{2+} → Fe²⁺
    同时作为兜底函数处理无法转换为OMML的简单化学式
    """
    # Unicode下标映射
    sub_map = {
        '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
        '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
        '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
        'a': 'ₐ', 'e': 'ₑ', 'h': 'ₕ', 'i': 'ᵢ', 'k': 'ₖ',
        'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'o': 'ₒ', 'p': 'ₚ',
        'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ', 'v': 'ᵥ', 'x': 'ₓ',
    }
    # Unicode上标映射
    sup_map = {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
        '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
        'n': 'ⁿ', 'i': 'ⁱ',
    }

    # 处理 _x 下标
    def replace_sub(match):
        chars = match.group(1)
        result = ''
        for c in chars:
            result += sub_map.get(c, c)
        return result
    text = re.sub(r'_\{([^}]+)\}', replace_sub, text)
    text = re.sub(r'_(\d|[+\-=()a-z])', lambda m: sub_map.get(m.group(1), m.group(1)), text)

    # 处理 ^{x} 上标
    def replace_sup(match):
        chars = match.group(1)
        result = ''
        for c in chars:
            result += sup_map.get(c, c)
        return result
    text = re.sub(r'\^\{([^}]+)\}', replace_sup, text)
    text = re.sub(r'\^(\d|[+\-=()a-z])', lambda m: sup_map.get(m.group(1), m.group(1)), text)

    return text


def _clean_latex_for_word(text: str) -> str:
    """
    清理LaTeX标记，转为Word可显示的文本
    复杂公式保留$...$标记(将转为OMML)，简单化学式转Unicode
    """
    # 先转换LaTeX符号
    text = _latex_to_unicode(text)

    # 拆分：$...$内的公式保留，外部文字转Unicode上下标
    parts = re.split(r'(\$[^$]+\$)', text)
    result_parts = []
    for part in parts:
        if part.startswith('$') and part.endswith('$'):
            result_parts.append(part)  # 保留公式标记
        else:
            result_parts.append(_parse_latex_subscript_superscript(part))

    return ''.join(result_parts)


def _add_formula_run(paragraph, latex_formula: str):
    """
    在Word段落中添加化学公式 run
    简单化学式使用Unicode上下标，复杂公式保留LaTeX标记
    """
    # 去掉$标记
    formula = latex_formula.strip('$')

    # 判断是否为简单化学式（只含元素符号、数字、上下标、+-、箭头等）
    is_simple_chemical = bool(re.match(
        r'^[A-Za-z0-9\s_{}\^+\-→←⇌→↑↓·()\[\]Δδθαβγλμ±°]+$',
        formula
    ))

    if is_simple_chemical:
        # 简单化学式：直接转Unicode上下标
        converted = _parse_latex_subscript_superscript(formula)
        run = paragraph.add_run(converted)
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        # 复杂公式：保留LaTeX格式，添加灰色背景提示
        run = paragraph.add_run(f"[{formula}]")
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(100, 100, 100)


def _add_mixed_text(paragraph, text: str, image_map: dict = None):
    """
    添加混合文本（含$...$公式标记和{{img:xxx}}图片占位符）到Word段落

    参数:
        paragraph: Word 段落对象
        text: 混合文本
        image_map: {img_id: 图片路径或bytes} 图片映射表
    """
    # 分割 $...$ 公式 和 {{img:xxx}} 图片占位符
    parts = re.split(r'(\$[^$]+\$)|(\{\{img:(\w+)\}\})', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('$') and part.endswith('$'):
            _add_formula_run(paragraph, part)
        elif part.startswith('{{img:') and image_map:
            img_id = part[5:-2]  # 提取 img_001
            img_path = image_map.get(img_id)
            if img_path:
                try:
                    run = paragraph.add_run()
                    # 控制图片宽度为 A4 页面的 60%（约 10cm）
                    run.add_picture(img_path, width=Cm(10))
                    paragraph.add_run('\n')  # 图片后换行
                except Exception:
                    # 图片加载失败，插入占位文字
                    run = paragraph.add_run(f'[{img_id}]')
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(150, 150, 150)
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def _add_question_images(doc, image_paths: list[str] | None):
    """Append question images as standalone Word paragraphs."""
    for image_path in image_paths or []:
        try:
            img_para = doc.add_paragraph()
            img_para.paragraph_format.left_indent = Cm(0.5)
            img_para.paragraph_format.space_before = Pt(3)
            img_para.paragraph_format.space_after = Pt(3)
            img_run = img_para.add_run()
            img_run.add_picture(image_path, width=Cm(10))
        except Exception:
            fallback_para = doc.add_paragraph()
            fallback_para.paragraph_format.left_indent = Cm(0.5)
            fallback_run = fallback_para.add_run("[image unavailable]")
            fallback_run.font.size = Pt(9)
            fallback_run.font.color.rgb = RGBColor(150, 150, 150)


def generate_test_paper_word(
    paper_title: str,
    paper_subtitle: str,
    total_score: int,
    exam_duration: int,
    questions: list[dict],
    image_map: dict = None,
) -> str:
    """
    生成化学试卷Word文档

    参数:
        questions: [{id, content, question_type, options, score, sort_order, answer, analysis}]

    返回:
        生成的.docx文件路径
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)   # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # ── 设置默认字体 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 试卷标题 ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(paper_title)
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    if paper_subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(paper_subtitle)
        sub_run.font.size = Pt(12)
        sub_run.font.name = '宋体'
        sub_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 试卷信息栏 ──
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    info_text = f"总分：{total_score}分    考试时间：{exam_duration}分钟"
    info_run = info_para.add_run(info_text)
    info_run.font.size = Pt(10.5)
    info_run.font.name = '宋体'
    info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ── 得分栏 ──
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    score_run = score_para.add_run("得分：__________")
    score_run.font.size = Pt(10.5)

    # 分隔线
    sep_para = doc.add_paragraph()
    sep_run = sep_para.add_run("─" * 60)
    sep_run.font.size = Pt(8)

    # ── 逐题添加 ──
    # 按题型分组
    type_order = {"choice": 0, "fill": 1, "experiment": 2, "calculation": 3, "short_answer": 4}
    type_names = {
        "choice": "选择题", "fill": "填空题", "experiment": "实验题",
        "calculation": "计算题", "short_answer": "简答题"
    }

    grouped = {}
    for q in questions:
        qt = q.get("question_type", "short_answer")
        if qt not in grouped:
            grouped[qt] = []
        grouped[qt].append(q)

    # 按题型顺序输出
    global_qno = 0
    for qt in sorted(grouped.keys(), key=lambda x: type_order.get(x, 99)):
        qs = grouped[qt]
        type_name = type_names.get(qt, "解答题")

        # 题型标题
        type_para = doc.add_paragraph()
        type_run = type_para.add_run(f"{type_name}（共{len(qs)}题）")
        type_run.font.size = Pt(14)
        type_run.font.bold = True
        type_run.font.name = '黑体'
        type_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for q in qs:
            global_qno += 1
            content = q.get("content", "")
            score = q.get("score", 5)

            # 题号 + 题目正文
            q_para = doc.add_paragraph()
            q_para.paragraph_format.space_before = Pt(6)
            q_para.paragraph_format.space_after = Pt(3)

            # 题号
            qno_run = q_para.add_run(f"{global_qno}. ")
            qno_run.font.size = Pt(11)
            qno_run.font.bold = True

            # 题目内容（支持化学公式 + 图片）
            _add_mixed_text(q_para, content, image_map)

            # 分值标注
            score_run = q_para.add_run(f"  （{score}分）")
            score_run.font.size = Pt(9)
            score_run.font.color.rgb = RGBColor(128, 128, 128)

            # 选择题选项
            if qt == "choice" and q.get("options"):
                for opt in q["options"]:
                    opt_label = opt.get("label", "")
                    opt_text = opt.get("text", "")
                    opt_para = doc.add_paragraph()
                    opt_para.paragraph_format.left_indent = Cm(1)
                    _add_mixed_text(opt_para, f"{opt_label}. {opt_text}", image_map)

            _add_question_images(doc, q.get("images"))

            # 答题空行（非选择题）
            if qt != "choice":
                blank_para = doc.add_paragraph()
                blank_para.paragraph_format.space_before = Pt(4)
                blank_para.paragraph_format.space_after = Pt(4)
                blank_run = blank_para.add_run("\n")
                blank_run.font.size = Pt(18)

    # ── 保存 ──
    import os
    export_dir = "./exports"
    os.makedirs(export_dir, exist_ok=True)
    filepath = os.path.join(export_dir, f"test_paper_{__import__('uuid').uuid4().hex[:8]}.docx")
    doc.save(filepath)
    return filepath


def generate_answer_sheet_word(
    paper_title: str,
    questions: list[dict],
    image_map: dict = None,
) -> str:
    """
    生成答案卷Word文档
    """
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(f"{paper_title} — 参考答案")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.name = '黑体'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    # 逐题
    for idx, q in enumerate(questions):
        qno = idx + 1

        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(6)

        qno_run = q_para.add_run(f"{qno}. ")
        qno_run.font.size = Pt(11)
        qno_run.font.bold = True

        answer = q.get("answer") or "暂无答案"
        _add_mixed_text(q_para, f"答案：{answer}", image_map)

        _add_question_images(doc, q.get("images"))

        analysis = q.get("analysis")
        if analysis:
            ana_para = doc.add_paragraph()
            ana_para.paragraph_format.left_indent = Cm(0.5)
            ana_run = ana_para.add_run("解析：")
            ana_run.font.size = Pt(10)
            ana_run.font.color.rgb = RGBColor(0, 100, 0)
            _add_mixed_text(ana_para, analysis, image_map)

    import os
    export_dir = "./exports"
    os.makedirs(export_dir, exist_ok=True)
    filepath = os.path.join(export_dir, f"answer_sheet_{__import__('uuid').uuid4().hex[:8]}.docx")
    doc.save(filepath)
    return filepath
